from __future__ import annotations

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


def _insert_paragraph_after(anchor: Paragraph, *, text: str, style) -> Paragraph:
    """Insert a new paragraph after `anchor` with optional style."""

    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    p = Paragraph(new_p, anchor._parent)
    if style is not None:
        try:
            p.style = style
        except Exception:
            # If style can't be applied, keep default
            pass
    p.add_run(text)
    return p


def _ensure_company_placeholder(doc: Document, *, role_title_placeholder: str, company_placeholder: str) -> bool:
    """Ensure `company_placeholder` paragraph exists right after `role_title_placeholder`.

    Returns True if a change was made.
    """

    # Flatten paragraphs (including tables) in document order.
    paras: list[Paragraph] = []
    paras.extend(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paras.extend(cell.paragraphs)

    for i, p in enumerate(paras):
        if (p.text or "").strip() != role_title_placeholder:
            continue

        # If the next paragraph already contains the company placeholder, do nothing.
        if i + 1 < len(paras) and company_placeholder in (paras[i + 1].text or ""):
            return False

        # Choose style from the next paragraph (typically the date line) if possible.
        style = paras[i + 1].style if i + 1 < len(paras) else p.style

        _insert_paragraph_after(p, text=company_placeholder, style=style)
        return True

    raise RuntimeError(f"Could not find role title placeholder paragraph: {role_title_placeholder!r}")


def main() -> int:
    path = "corey_resume_template.docx"
    doc = Document(path)

    changed = False
    changed |= _ensure_company_placeholder(
        doc,
        role_title_placeholder="<< Role-Title-1 >>",
        company_placeholder="<< Company-1 >>",
    )
    changed |= _ensure_company_placeholder(
        doc,
        role_title_placeholder="<< Role-Title-3 >>",
        company_placeholder="<< Company-3 >>",
    )
    changed |= _ensure_company_placeholder(
        doc,
        role_title_placeholder="<< Role-Title-4 >>",
        company_placeholder="<< Company-4 >>",
    )

    if changed:
        doc.save(path)
        print("UPDATED", path)
    else:
        print("NO_CHANGES", path)

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
