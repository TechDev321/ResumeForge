"""Patch generated resume DOCX titles for consistency.

Usage (Windows):
  python tools/patch_docx_titles.py --in in.docx --out out.docx

This is a small utility for one-off patching of already-generated resumes.
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document


def _replace_in_paragraph(par, replacements: list[tuple[str, str]]) -> None:
    # Replace within runs to preserve formatting as much as possible.
    for run in par.runs:
        text = run.text
        if not text:
            continue
        for old, new in replacements:
            if old in text:
                text = text.replace(old, new)
        run.text = text


def patch_docx(*, in_path: Path, out_path: Path) -> None:
    doc = Document(str(in_path))

    replacements = [
        ("Junior Software Engineer", "Software Engineer"),
        ("Junior software engineer", "Software engineer"),
    ]

    for p in doc.paragraphs:
        _replace_in_paragraph(p, replacements)

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, replacements)

    doc.save(str(out_path))


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--in", dest="in_path", required=True)
    ap.add_argument("--out", dest="out_path", required=True)
    args = ap.parse_args()

    in_path = Path(args.in_path)
    out_path = Path(args.out_path)
    patch_docx(in_path=in_path, out_path=out_path)
    print(str(out_path))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
