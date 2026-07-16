from __future__ import annotations

import os
from io import BytesIO
from pathlib import Path

from docx import Document

from .cover_letter_template import create_cover_letter_template
from .docx_utils import iter_paragraphs
from .openai_cover_letter import CoverLetterGenerationError, generate_cover_letter_replacements
from .openai_question import generate_interview_answer
from .openai_resume import ResumeGenerationError, generate_replacements
from .template_engine import apply_replacements, extract_placeholders


BACKEND_DIR = Path(__file__).resolve().parent.parent

DEFAULT_TEMPLATE = "corey_resume_template(Nielsen_Disney_TCS).docx"
DEFAULT_PROMPT = "Resume Prompt.txt"
DEFAULT_COVER_LETTER_TEMPLATE = "corey_cover_letter_template.docx"
DEFAULT_COVER_LETTER_PROMPT = "Cover Letter Prompt.txt"
DEFAULT_QUESTION_PROMPT = "Question_prompt.txt"


def strip_tag_block(text: str, tag: str) -> str:
    """Remove a <TAG>...</TAG> block if present (case-insensitive)."""

    lower = text.lower()
    start_tag = f"<{tag.lower()}>"
    end_tag = f"</{tag.lower()}>"
    if start_tag not in lower or end_tag not in lower:
        return text
    start = lower.index(start_tag)
    end = lower.index(end_tag)
    return text[:start] + text[end + len(end_tag) :]


def resolve_path(path: str | Path) -> Path:
    p = Path(path)
    if p.is_absolute():
        return p
    return BACKEND_DIR / p


def load_prompt_instructions(prompt_path: Path) -> str:
    prompt_rules = prompt_path.read_text(encoding="utf-8")
    # Resume Prompt.txt may contain example <JD> and <resume> blocks.
    # Strip those so only instructions go to the model for new jobs.
    prompt_instructions = strip_tag_block(prompt_rules, "JD")
    prompt_instructions = strip_tag_block(prompt_instructions, "resume")
    return prompt_instructions


def generate_resume_docx(
    *,
    jd_text: str,
    api_key: str,
    model: str | None = None,
    template_path: str | Path | None = None,
    prompt_path: str | Path | None = None,
    cache_dir: Path | None = None,
) -> tuple[str, bytes]:
    """Generate a filled resume. Returns (file_name, docx_bytes)."""

    jd = (jd_text or "").strip()
    if not jd:
        raise ValueError("Job description is empty")

    model = model or os.getenv("OPENAI_MODEL", "gpt-4.1-mini")
    template = resolve_path(template_path or os.getenv("RESUME_TEMPLATE", DEFAULT_TEMPLATE))
    prompt = resolve_path(prompt_path or os.getenv("RESUME_PROMPT", DEFAULT_PROMPT))

    if not template.is_file():
        raise FileNotFoundError(
            f"Resume template not found: {template}. "
            "Place the .docx in Backend/ or set RESUME_TEMPLATE."
        )
    if not prompt.is_file():
        raise FileNotFoundError(f"Resume prompt not found: {prompt}")

    prompt_instructions = load_prompt_instructions(prompt)

    doc_tmp = Document(str(template))
    resume_template_text = "\n".join(p.text for p in iter_paragraphs(doc_tmp))

    doc = Document(str(template))
    placeholders = extract_placeholders(doc)
    if not placeholders:
        raise ResumeGenerationError("No placeholders found in template.")

    file_name, replacements = generate_replacements(
        api_key=api_key,
        model=model,
        prompt_rules=prompt_instructions,
        jd_text=jd,
        resume_template_text=resume_template_text,
        placeholders=placeholders,
        cache_dir=cache_dir,
    )

    apply_replacements(doc, replacements)

    buffer = BytesIO()
    doc.save(buffer)
    return file_name, buffer.getvalue()


def read_docx_text(docx_bytes: bytes) -> str:
    """Extract plain paragraph text from a .docx byte payload."""

    doc = Document(BytesIO(docx_bytes))
    return "\n".join(p.text for p in iter_paragraphs(doc) if (p.text or "").strip())


def generate_cover_letter_docx(
    *,
    jd_text: str,
    resume_bytes: bytes,
    api_key: str,
    model: str | None = None,
    template_path: str | Path | None = None,
    prompt_path: str | Path | None = None,
    cache_dir: Path | None = None,
) -> tuple[str, bytes]:
    """Generate a filled cover letter from JD + resume .docx bytes.

    Returns (file_name, docx_bytes).
    """

    jd = (jd_text or "").strip()
    if not jd:
        raise ValueError("Job description is empty")
    if not resume_bytes:
        raise ValueError("Resume file is empty")

    resume_text = read_docx_text(resume_bytes)
    if not resume_text.strip():
        raise ValueError("Could not extract text from the resume .docx")

    model = model or os.getenv("OPENAI_MODEL", "gpt-4.1-mini")
    template = resolve_path(
        template_path or os.getenv("COVER_LETTER_TEMPLATE", DEFAULT_COVER_LETTER_TEMPLATE)
    )
    prompt = resolve_path(
        prompt_path or os.getenv("COVER_LETTER_PROMPT", DEFAULT_COVER_LETTER_PROMPT)
    )

    if not template.is_file():
        create_cover_letter_template(template)
    if not prompt.is_file():
        raise FileNotFoundError(f"Cover letter prompt not found: {prompt}")

    prompt_rules = prompt.read_text(encoding="utf-8")

    doc = Document(str(template))
    placeholders = extract_placeholders(doc)
    if not placeholders:
        raise CoverLetterGenerationError("No placeholders found in cover letter template.")

    file_name, replacements = generate_cover_letter_replacements(
        api_key=api_key,
        model=model,
        prompt_rules=prompt_rules,
        jd_text=jd,
        resume_text=resume_text,
        placeholders=placeholders,
        cache_dir=cache_dir,
    )

    apply_replacements(doc, replacements)

    buffer = BytesIO()
    doc.save(buffer)
    return file_name, buffer.getvalue()


def generate_question_answer(
    *,
    jd_text: str,
    resume_bytes: bytes,
    question_text: str,
    api_key: str,
    model: str | None = None,
    prompt_path: str | Path | None = None,
    cache_dir: Path | None = None,
) -> str:
    """Generate an interview answer from JD + resume .docx + question."""

    jd = (jd_text or "").strip()
    if not jd:
        raise ValueError("Job description is empty")
    if not resume_bytes:
        raise ValueError("Resume file is empty")
    question = (question_text or "").strip()
    if not question:
        raise ValueError("Question is empty")

    resume_text = read_docx_text(resume_bytes)
    if not resume_text.strip():
        raise ValueError("Could not extract text from the resume .docx")

    model = model or os.getenv("OPENAI_MODEL", "gpt-4.1-mini")
    prompt = resolve_path(prompt_path or os.getenv("QUESTION_PROMPT", DEFAULT_QUESTION_PROMPT))
    if not prompt.is_file():
        raise FileNotFoundError(f"Question prompt not found: {prompt}")

    prompt_rules = prompt.read_text(encoding="utf-8")

    return generate_interview_answer(
        api_key=api_key,
        model=model,
        prompt_rules=prompt_rules,
        jd_text=jd,
        resume_text=resume_text,
        question_text=question,
        cache_dir=cache_dir,
    )
