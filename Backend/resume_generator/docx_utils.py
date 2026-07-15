from __future__ import annotations

from dataclasses import dataclass
from typing import Iterator

from copy import deepcopy
from docx.text.run import Run


def iter_paragraphs(document) -> Iterator:
    """Yield paragraphs from a python-docx Document, including those inside tables."""

    # Normal body paragraphs
    for p in document.paragraphs:
        yield p

    # Paragraphs inside tables (if any)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


@dataclass(frozen=True)
class TextRun:
    text: str
    bold: bool = False


def parse_bold_markdown(text: str, marker: str = "**") -> list[TextRun]:
    """Parse a minimal markdown-like bold marker (**bold**) into runs.

    This is intentionally tiny/safe:
    - Only supports a single marker string (default **)
    - No nesting/escaping
    - Unmatched markers are treated as literal text
    """

    if not text:
        return [TextRun("")]

    parts: list[TextRun] = []
    i = 0
    bold = False
    m = marker
    ml = len(m)

    buf: list[str] = []

    def flush():
        nonlocal buf
        if buf:
            parts.append(TextRun("".join(buf), bold=bold))
            buf = []

    while i < len(text):
        if text.startswith(m, i):
            flush()
            bold = not bold
            i += ml
            continue
        buf.append(text[i])
        i += 1

    flush()

    # If we ended "inside bold" due to unmatched marker, revert by merging runs
    if bold:
        # join everything as literal
        return [TextRun(text, bold=False)]

    return parts


def replace_paragraph_text_with_runs(paragraph, runs: list[TextRun]) -> None:
    """Replace paragraph runs while preserving paragraph-level formatting."""

    # Capture run-level formatting from the first run, if present, so we can
    # keep the template's fonts/sizes/colors as closely as python-docx allows.
    base = paragraph.runs[0] if paragraph.runs else None
    base_style = getattr(base, "style", None) if base is not None else None
    base_font = getattr(base, "font", None) if base is not None else None
    base_font_name = getattr(base_font, "name", None) if base_font is not None else None
    base_font_size = getattr(base_font, "size", None) if base_font is not None else None
    base_font_color = getattr(getattr(base_font, "color", None), "rgb", None) if base_font is not None else None
    base_italic = getattr(base, "italic", None) if base is not None else None
    base_underline = getattr(base, "underline", None) if base is not None else None

    # Clear existing runs
    # python-docx has no public API; this is the standard approach.
    p = paragraph._p
    for r in list(p.r_lst):
        p.remove(r)

    for tr in runs:
        r = paragraph.add_run(tr.text)

        # Apply base formatting first
        if base_style is not None:
            r.style = base_style
        if base_font_name is not None:
            r.font.name = base_font_name
        if base_font_size is not None:
            r.font.size = base_font_size
        if base_font_color is not None:
            r.font.color.rgb = base_font_color
        if base_italic is not None:
            r.italic = base_italic
        if base_underline is not None:
            r.underline = base_underline

        # Then apply bold markers
        if tr.bold:
            r.bold = True


def _clone_run_formatting(*, src: Run, dst: Run) -> None:
    """Clone run formatting (the <w:rPr> element) from src to dst.

    python-docx doesn't provide a full-fidelity public API for copying complex
    run formatting (character styles, theme fonts, spacing, etc.). For maximum
    preservation of the original template formatting we copy the underlying
    XML run properties.
    """

    try:
        src_rpr = src._r.get_or_add_rPr()
    except Exception:
        src_rpr = None

    if src_rpr is None:
        return

    # Remove any existing rPr on destination run and replace it.
    dst_r = dst._r
    dst_rpr = dst_r.rPr
    if dst_rpr is not None:
        dst_r.remove(dst_rpr)
    dst_r.append(deepcopy(src_rpr))


def replace_placeholders_in_paragraph_runs(paragraph, replacements: dict[str, str]) -> bool:
    """Replace placeholders *within runs* to preserve all run-level styles.

    Returns True if it handled any replacements, else False.

    This intentionally only supports placeholders that exist wholly inside a
    single run. If Word split a placeholder across multiple runs, this function
    returns False so the caller can fall back to a paragraph-wide rewrite.
    """

    if not paragraph.runs:
        return False

    did_any = False

    # We repeatedly rebuild the paragraph's full run text and apply the first
    # found replacement. This is slightly less "clever" than computing all
    # match locations up-front, but it's much safer because Word templates can
    # split placeholder tokens across multiple runs.
    while True:
        full = "".join(r.text or "" for r in paragraph.runs)
        if "<<" not in full:
            break

        # Find the earliest placeholder occurrence in this paragraph.
        found_key = None
        found_at = None
        for k in replacements.keys():
            pos = full.find(k)
            if pos == -1:
                continue
            if found_at is None or pos < found_at:
                found_key = k
                found_at = pos

        if found_key is None or found_at is None:
            break

        repl_text = replacements.get(found_key, "")
        start = found_at
        end = found_at + len(found_key)

        # Locate start/end run boundaries.
        cursor = 0
        start_run_idx = None
        start_off = None
        end_run_idx = None
        end_off = None

        runs = list(paragraph.runs)
        for i, r in enumerate(runs):
            t = r.text or ""
            next_cursor = cursor + len(t)

            if start_run_idx is None and start < next_cursor:
                start_run_idx = i
                start_off = start - cursor

            if end_run_idx is None and end <= next_cursor:
                end_run_idx = i
                end_off = end - cursor
                break

            cursor = next_cursor

        if start_run_idx is None or end_run_idx is None or start_off is None or end_off is None:
            # Defensive: shouldn't happen.
            break

        start_run = runs[start_run_idx]
        end_run = runs[end_run_idx]

        prefix = (start_run.text or "")[:start_off]
        suffix = (end_run.text or "")[end_off:]

        parent = start_run._r.getparent()
        insert_at = parent.index(start_run._r)

        # Remove the runs that are fully/partially covered by the placeholder.
        for r in runs[start_run_idx : end_run_idx + 1]:
            try:
                r._r.getparent().remove(r._r)
            except Exception:
                pass

        # Insert prefix (keeps formatting of the run where placeholder started)
        if prefix:
            new_r = paragraph.add_run(prefix)
            _clone_run_formatting(src=start_run, dst=new_r)
            parent.remove(new_r._r)
            parent.insert(insert_at, new_r._r)
            insert_at += 1

        # Insert replacement runs (clone start-run formatting)
        for tr in parse_bold_markdown(repl_text):
            if tr.text == "":
                continue
            new_r = paragraph.add_run(tr.text)
            _clone_run_formatting(src=start_run, dst=new_r)
            if tr.bold:
                new_r.bold = True
            parent.remove(new_r._r)
            parent.insert(insert_at, new_r._r)
            insert_at += 1

        # Insert suffix (keeps formatting of the run where placeholder ended)
        if suffix:
            new_r = paragraph.add_run(suffix)
            _clone_run_formatting(src=end_run, dst=new_r)
            parent.remove(new_r._r)
            parent.insert(insert_at, new_r._r)
            insert_at += 1

        did_any = True

    return did_any
